import os
import requests

# All api imports
from fastapi import FastAPI, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

# All RAG model imports
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

# Misc. imports for reading files
from docx import Document as DocxDocument
import fitz  # PyMuPDF

app = FastAPI()

# Enable CORS for local frontend or Open WebUI
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===================== Text Extraction =====================
def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    return "\n".join(page.get_text() for page in doc)

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def load_documents(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(file_path)
        elif filename.endswith(".txt"):
            text = extract_text_from_txt(file_path)
        else:
            continue
        docs.append(Document(page_content=text, metadata={"source": filename}))
    return docs

# ===================== Text Splitting =====================
def split_documents(docs):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_documents(docs)

# ===================== Embeddings =====================
def build_vectorstore(docs):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return FAISS.from_documents(docs, embeddings)

# ===================== Local LLM Query via Ollama =====================
def query_ollama(prompt, model="llama2-uncensored"):
    response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": model,
            "prompt": prompt,
            "stream": False
        }
    )
    return response.json()["response"]

# ===================== RAG Answer Generation =====================
def ask_question(query, vectorstore):
    retriever = vectorstore.as_retriever()
    relevant_docs = retriever.get_relevant_documents(query)
    context = "\n\n".join(doc.page_content for doc in relevant_docs)
    prompt = f"Answer the question based on the context below:\n\n{context}\n\nQuestion: {query}\nAnswer:"
    answer = query_ollama(prompt)
    return answer, [doc.metadata["source"] for doc in relevant_docs]

# ===================== FastAPI Endpoint =====================
@app.post("/rag-query/")
async def rag_query(folder_path: str = Form(...), query: str = Form(...)):
    try:
        raw_docs = load_documents(folder_path)
        split_docs = split_documents(raw_docs)
        vectorstore = build_vectorstore(split_docs)
        answer, sources = ask_question(query, vectorstore)
        return JSONResponse({"answer": answer, "sources": sources})
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

# ===================== Entry Point (for manual testing) =====================
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("rag_fastapi_app:app", host="0.0.0.0", port=8000, reload=True)